from docx import Document
import markdown
import re
import tempfile
import os
import io

class DocumentConverter:
    @staticmethod
    def docx_to_markdown(file_content: bytes) -> str:
        # Create a temporary file with .docx extension
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            try:
                # Write the content to the temp file
                temp_file.write(file_content)
                temp_file.flush()
                
                # Open the document with python-docx
                doc = Document(temp_file.name)
                
                # Convert to markdown, preserving structure
                markdown_content = []
                
                # Process paragraphs
                for paragraph in doc.paragraphs:
                    # Skip empty paragraphs
                    if not paragraph.text.strip():
                        continue
                        
                    # Handle different paragraph styles
                    if paragraph.style.name.startswith('Heading'):
                        # Add appropriate number of '#' for heading level
                        level = int(paragraph.style.name[-1])
                        markdown_content.append(f"{'#' * level} {paragraph.text}")
                    else:
                        markdown_content.append(paragraph.text)
                
                # Process tables if any
                for table in doc.tables:
                    rows = []
                    for row in table.rows:
                        cells = [cell.text for cell in row.cells]
                        rows.append('| ' + ' | '.join(cells) + ' |')
                    
                    if rows:
                        # Add header separator
                        header_sep = '| ' + ' | '.join(['---'] * len(table.columns)) + ' |'
                        rows.insert(1, header_sep)
                        markdown_content.extend(rows)
                
                return "\n\n".join(markdown_content)
                
            finally:
                # Clean up: close and remove the temporary file
                temp_file.close()
                os.unlink(temp_file.name)
    
    @staticmethod
    def find_blanks(markdown_text: str) -> list[str]:
        # Find text within square brackets or sequences of underscores
        blank_patterns = [
            r'\[([^\]]+)\]',  # Text within square brackets
            r'_{3,}',         # Three or more underscores
        ]
        
        blanks = []
        for pattern in blank_patterns:
            matches = re.finditer(pattern, markdown_text)
            for match in matches:
                # For bracketed text, get the content within brackets
                if '[' in pattern:
                    blanks.append(match.group(0))  # Include the brackets
                else:
                    # For underscores, get the whole match
                    blanks.append(match.group(0))
        
        return blanks
    
    def markdown_to_docx(self, markdown_text: str) -> bytes:
        doc = Document()
        
        # Split into paragraphs
        paragraphs = markdown_text.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                # Handle headers
                if para_text.startswith('#'):
                    level = len(para_text.split()[0])  # Count #'s
                    text = para_text.lstrip('#').strip()
                    doc.add_heading(text, level=level)
                else:
                    doc.add_paragraph(para_text)
        
        # Save to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes.read()